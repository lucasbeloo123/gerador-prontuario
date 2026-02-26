from flask import Flask, jsonify, send_from_directory
from docx import Document
import uuid
import os

app = Flask(__name__)

UPLOAD_FOLDER = "files"
os.makedirs(UPLOAD_FOLDER, exist_ok=True)

@app.route("/gerar-prontuario", methods=["POST"])
def gerar_prontuario():

    doc = Document()
    doc.add_heading('PRONTUÁRIO MÉDICO', 0)

    campos = [
        "IDENTIFICAÇÃO DO PACIENTE",
        "Nome:",
        "Data de Nascimento:",
        "Sexo:",
        "CPF:",
        "Telefone:",
        "Endereço:",
        "",
        "ANAMNESE",
        "Queixa Principal:",
        "História da Doença Atual:",
        "",
        "ANTECEDENTES PESSOAIS",
        "Doenças prévias:",
        "Cirurgias:",
        "Alergias:",
        "Medicações em uso:",
        "",
        "ANTECEDENTES FAMILIARES",
        "",
        "EXAME FÍSICO",
        "",
        "HIPÓTESE DIAGNÓSTICA",
        "",
        "PLANO TERAPÊUTICO",
        "",
        "PRESCRIÇÃO",
        "",
        "EVOLUÇÃO",
        "",
        "OBSERVAÇÕES ADICIONAIS"
    ]

    for campo in campos:
        doc.add_paragraph(campo)

    nome_arquivo = f"{uuid.uuid4()}.docx"
    caminho = os.path.join(UPLOAD_FOLDER, nome_arquivo)
    doc.save(caminho)

    file_url = f"https://gerador-prontuario.onrender.com"

    return jsonify({
        "file_url": file_url
    })


@app.route("/files/<filename>")
def baixar_arquivo(filename):
    return send_from_directory(UPLOAD_FOLDER, filename, as_attachment=True)


if __name__ == "__main__":
    app.run(host="0.0.0.0", port=10000)
